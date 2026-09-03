import os
import sys
import tempfile
import unittest
from PyQt6.QtWidgets import QApplication, QTreeWidgetItem
from PyQt6.QtCore import Qt

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from views.main_window import MainWindow
from controllers.main_controller import MainController
from docx import Document

app = QApplication.instance()
if not app:
    app = QApplication(sys.argv)


class TestExportFormats(unittest.TestCase):
    """測試 ExportController 支援的多格式匯出 (docx, txt, md, epub)。"""

    def setUp(self):
        self.temp_dir = tempfile.TemporaryDirectory()
        self.view = MainWindow()
        self.mc = MainController(self.view)

        # 建立測試節點 (包含 Markdown 標記)
        self.item1 = self.mc.tree.create_item("第一章 啟程", is_folder=False, content="這是**第一章**的內文，他*輕聲說道*。\n---\n> 這是一封密信。")
        self.item2 = self.mc.tree.create_item("第二章 冒險", is_folder=False, content="這是第二章的內文，遭遇了~~強大的對手~~強敵。")
        self.files_list = [self.item1, self.item2]

    def tearDown(self):
        if hasattr(self.mc, 'writing_timer') and self.mc.writing_timer.isActive():
            self.mc.writing_timer.stop()
        if hasattr(self.mc, 'auto_save_timer') and self.mc.auto_save_timer.isActive():
            self.mc.auto_save_timer.stop()
        self.view.close()
        self.temp_dir.cleanup()

    def test_export_docx(self):
        docx_path = os.path.join(self.temp_dir.name, "test.docx")
        self.mc.export_controller._save_as_docx(self.files_list, docx_path, include_title=True)
        self.assertTrue(os.path.exists(docx_path))
        doc = Document(docx_path)
        self.assertTrue(len(doc.paragraphs) >= 4)
        self.assertEqual(doc.paragraphs[0].text, "第一章 啟程")
        # 驗證是否有粗體 Run
        has_bold_run = any(r.bold and r.text == "第一章" for p in doc.paragraphs for r in p.runs)
        self.assertTrue(has_bold_run)

    def test_export_txt(self):
        txt_path = os.path.join(self.temp_dir.name, "test.txt")
        self.mc.export_controller._save_as_txt(self.files_list, txt_path, include_title=True)
        self.assertTrue(os.path.exists(txt_path))
        with open(txt_path, "r", encoding="utf-8") as f:
            content = f.read()
        self.assertIn("【第一章 啟程】", content)
        # 驗證 Markdown 標記已乾淨去除並保留純文字
        self.assertIn("這是第一章的內文，他輕聲說道。", content)
        self.assertNotIn("**第一章**", content)
        self.assertIn("【第二章 冒險】", content)

    def test_export_md(self):
        md_path = os.path.join(self.temp_dir.name, "test.md")
        self.mc.export_controller._save_as_md(self.files_list, md_path, include_title=True)
        self.assertTrue(os.path.exists(md_path))
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()
        self.assertIn("# 第一章 啟程", content)
        self.assertIn("這是**第一章**的內文", content)
        self.assertIn("# 第二章 冒險", content)

    def test_export_epub(self):
        epub_path = os.path.join(self.temp_dir.name, "test.epub")
        self.mc.export_controller._save_as_epub(self.files_list, epub_path, include_title=True, book_title="測試小說")
        self.assertTrue(os.path.exists(epub_path))
        self.assertTrue(os.path.getsize(epub_path) > 500)

    def test_export_default_dir_follows_storage_path(self):
        """測試匯出預設目錄正確跟隨 mc.get_export_dir()。"""
        export_dir = self.mc.get_export_dir()
        self.assertTrue(os.path.exists(export_dir))
        self.assertEqual(os.path.basename(export_dir), "Export")
        # 測試變更自訂存檔路徑後，匯出目錄連動改變
        custom_storage = os.path.join(self.temp_dir.name, "CustomStoryRoot")
        self.mc.app_settings["storage_path"] = custom_storage
        new_export_dir = self.mc.get_export_dir()
        self.assertEqual(new_export_dir, os.path.join(custom_storage, "Export"))


if __name__ == "__main__":
    unittest.main()

