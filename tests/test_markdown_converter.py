import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pytest
from utils.markdown_converter import MarkdownConverter


def test_parse_inline_tokens_plain():
    tokens = MarkdownConverter.parse_inline_tokens("這是一段純文字。")
    assert len(tokens) == 1
    assert tokens[0]["text"] == "這是一段純文字。"
    assert not tokens[0]["bold"]
    assert not tokens[0]["italic"]


def test_parse_inline_tokens_mixed():
    text = "主角說道：**「別過來！」**他心中*暗自震驚*，甚至~~有些猶豫~~與`程式碼`。"
    tokens = MarkdownConverter.parse_inline_tokens(text)
    
    # 驗證 tokens 解析
    bold_tokens = [t for t in tokens if t["bold"] and not t["italic"]]
    assert len(bold_tokens) == 1
    assert bold_tokens[0]["text"] == "「別過來！」"

    italic_tokens = [t for t in tokens if t["italic"] and not t["bold"]]
    assert len(italic_tokens) == 1
    assert italic_tokens[0]["text"] == "暗自震驚"

    strike_tokens = [t for t in tokens if t["strike"]]
    assert len(strike_tokens) == 1
    assert strike_tokens[0]["text"] == "有些猶豫"

    code_tokens = [t for t in tokens if t["code"]]
    assert len(code_tokens) == 1
    assert code_tokens[0]["text"] == "程式碼"


def test_to_plain_text():
    md = "# 第一章 風起\n\n這是**重點**，他*輕聲說*。\n---\n> 這是一段信件。"
    plain = MarkdownConverter.to_plain_text(md, auto_indent=True)
    lines = plain.split('\n')
    
    assert lines[0] == "第一章 風起"
    assert lines[1] == ""
    assert "這是重點，他輕聲說。" in lines[2]
    assert lines[2].startswith("　　")
    assert lines[3] == "――――――――――――――――――――"
    assert "這是一段信件。" in lines[4]


def test_to_html_paragraphs():
    md = "# 標題\n**粗體**與*斜體*\n---\n> 引用內容"
    html_p = MarkdownConverter.to_html_paragraphs(md)
    
    assert html_p[0] == "<h1>標題</h1>"
    assert html_p[1] == "<p><strong>粗體</strong>與<em>斜體</em></p>"
    assert html_p[2] == "<hr/>"
    assert html_p[3] == "<blockquote><p>引用內容</p></blockquote>"


def test_to_html_paragraphs_empty_lines():
    # 測試單個空行僅為段落分隔，不插入多餘空白段
    md_single_gap = "第一段\n\n第二段"
    html_p = MarkdownConverter.to_html_paragraphs(md_single_gap)
    assert len(html_p) == 2
    assert html_p[0] == "<p>第一段</p>"
    assert html_p[1] == "<p>第二段</p>"

    # 測試連續 2 個以上空行保留場景過場
    md_double_gap = "第一段\n\n\n第二段"
    html_p2 = MarkdownConverter.to_html_paragraphs(md_double_gap)
    assert len(html_p2) == 3
    assert html_p2[0] == "<p>第一段</p>"
    assert html_p2[1] == '<p class="scene-break"><br/></p>'
    assert html_p2[2] == "<p>第二段</p>"


def test_render_to_docx():
    from docx import Document
    doc = Document()
    md = "# 第一章 試煉\n\n這是一個**關鍵線索**與*低語*。\n\n---\n> 來自遠方的信。"
    MarkdownConverter.render_to_docx(md, doc, "Iansui")
    
    # 檢查 paragraphs 數量
    assert len(doc.paragraphs) >= 4
    # 檢查第一段標題
    assert doc.paragraphs[0].text == "第一章 試煉"
    assert doc.paragraphs[0].runs[0].bold is True
    # 檢查第二段內文 runs
    p2 = doc.paragraphs[2]  # 空行在 idx 1
    bold_runs = [r for r in p2.runs if r.bold]
    assert len(bold_runs) == 1
    assert bold_runs[0].text == "關鍵線索"


def test_markdown_to_html_empty_line_style():
    from utils.markdown_utils import markdown_to_html, document_to_markdown
    from PyQt6.QtWidgets import QApplication, QTextEdit
    
    app = QApplication.instance() or QApplication([])
    
    # 測試單個空行與連續空行生成的 HTML 帶有 -qt-paragraph-type:empty 標記
    md = "第一段\n\n第二段\n\n\n第三段"
    html = markdown_to_html(md)
    assert '<p style="-qt-paragraph-type:empty;"><br></p>' in html
    
    # 測試載入至 QTextEdit 時，空行高度維持單行高度 (16px)，不產生雙倍行高
    edit = QTextEdit()
    edit.resize(500, 500)
    edit.show()
    edit.setHtml(html)
    app.processEvents()
    
    b_empty = edit.document().findBlockByNumber(1)
    height = edit.document().documentLayout().blockBoundingRect(b_empty).height()
    assert height <= 20.0, f"空行高度過大: {height}px，預期為單行字高 (約16px)"
    
    # 測試序列化回 Markdown 完整無損
    roundtrip_md = document_to_markdown(edit.document())
    assert roundtrip_md == md
    edit.close()

