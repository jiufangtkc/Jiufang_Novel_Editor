import unittest
import os
import tempfile
from services.import_service import ImportService, ImportOptions
from models.models import ChapterNode

class TestImportService(unittest.TestCase):
    def test_novel_regex_parsing(self):
        sample_text = """
這是全書的引子，介紹了江湖的背景。

第一卷 風雲際會
本卷說明天下大亂。

第一章 初出茅廬
少年李逍遙走出了客棧。
外面下著大雨。

第二章 巧遇貴人
李逍遙在路上遇到了一位白衣女子。

第二卷 潛龍在淵
第三章 京城風雲
京城戒備森嚴。
"""
        options = ImportOptions(mode="novel_regex")
        nodes = ImportService.parse_text(sample_text, "測試作品", options)

        # 預期：引子 (file), 第一卷 (folder -> 第一章, 第二章), 第二卷 (folder -> 第三章)
        self.assertEqual(len(nodes), 3)
        
        # 1. 序言/引子
        self.assertEqual(nodes[0].name, "序言 / 前言")
        self.assertEqual(nodes[0].node_type, "file")
        self.assertIn("這是全書的引子", nodes[0].content)

        # 2. 第一卷
        vol1 = nodes[1]
        self.assertEqual(vol1.name, "第一卷 風雲際會")
        self.assertEqual(vol1.node_type, "folder")
        # 包含：卷首導言、第一章、第二章 共 3 個節點
        self.assertEqual(len(vol1.children), 3)
        self.assertEqual(vol1.children[0].name, "序言 / 前言")
        self.assertIn("本卷說明天下大亂", vol1.children[0].content)
        self.assertEqual(vol1.children[1].name, "第一章 初出茅廬")
        self.assertEqual(vol1.children[1].node_type, "file")
        self.assertIn("少年李逍遙走出了客棧", vol1.children[1].content)
        self.assertEqual(vol1.children[2].name, "第二章 巧遇貴人")

        # 3. 第二卷
        vol2 = nodes[2]
        self.assertEqual(vol2.name, "第二卷 潛龍在淵")
        self.assertEqual(vol2.node_type, "folder")
        self.assertEqual(len(vol2.children), 1)
        self.assertEqual(vol2.children[0].name, "第三章 京城風雲")

    def test_scene_split(self):
        sample_text = """
第一章 決戰紫禁城
月圓之夜，紫禁之巔。
***
劍光一閃，勝負已分。
***
葉孤城長嘆一聲。
"""
        options = ImportOptions(mode="novel_regex", enable_scene_split=True)
        nodes = ImportService.parse_text(sample_text, "測試作品", options)

        self.assertEqual(len(nodes), 1)
        chap = nodes[0]
        self.assertEqual(chap.name, "第一章 決戰紫禁城")
        self.assertEqual(chap.node_type, "file")
        self.assertIn("月圓之夜", chap.content)
        self.assertEqual(len(chap.children), 2)
        self.assertEqual(chap.children[0].name, "場景 1")
        self.assertEqual(chap.children[0].node_type, "scene")
        self.assertIn("劍光一閃", chap.children[0].content)
        self.assertEqual(chap.children[1].name, "場景 2")
        self.assertIn("葉孤城長嘆", chap.children[1].content)

    def test_markdown_parsing(self):
        sample_md = """
# 卷一 天下大勢
## 第一章 出山
這是出山的內容。
### 第一節 告別師門
師父叮囑多加小心。
## 第二章 渡河
河面波濤洶湧。
"""
        nodes = ImportService.parse_markdown(sample_md, "MD作品")
        self.assertEqual(len(nodes), 1)
        vol1 = nodes[0]
        self.assertEqual(vol1.name, "卷一 天下大勢")
        self.assertEqual(vol1.node_type, "folder")
        self.assertEqual(len(vol1.children), 2)

        chap1 = vol1.children[0]
        self.assertEqual(chap1.name, "第一章 出山")
        self.assertIn("這是出山的內容", chap1.content)
        self.assertEqual(len(chap1.children), 1)
        self.assertEqual(chap1.children[0].name, "第一節 告別師門")
        self.assertEqual(chap1.children[0].node_type, "scene")

        chap2 = vol1.children[1]
        self.assertEqual(chap2.name, "第二章 渡河")

    def test_single_chapter_mode(self):
        text = "這是一篇極短篇小說，不分章節。"
        options = ImportOptions(mode="single_chapter")
        nodes = ImportService.parse_text(text, "短篇傳奇", options)
        self.assertEqual(len(nodes), 1)
        self.assertEqual(nodes[0].name, "短篇傳奇")
        self.assertEqual(nodes[0].content, text)

    def test_encoding_detection(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            # 測試 UTF-8
            utf8_path = os.path.join(tmpdir, "test_utf8.txt")
            with open(utf8_path, "w", encoding="utf-8") as f:
                f.write("第一章 繁體中文測試\n內文內容")
            enc = ImportService.detect_encoding(utf8_path)
            self.assertEqual(enc, "utf-8")

            # 測試 Big5 / CP950
            cp950_path = os.path.join(tmpdir, "test_cp950.txt")
            with open(cp950_path, "w", encoding="cp950") as f:
                f.write("第一章 測試繁體中文。\n這是一個在台灣使用繁體編碼的測試檔案，包含了許多常用字詞與標點符號！")
            enc_cp950 = ImportService.detect_encoding(cp950_path)
            self.assertEqual(enc_cp950, "cp950")
            content = ImportService.read_file_content(cp950_path)
            self.assertIn("第一章", content)
            self.assertIn("繁體中文", content)

    def test_docx_parsing(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("未安裝 python-docx")

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "novel.docx")
            doc = Document()
            doc.add_heading("第一卷 風起雲湧", level=1)
            doc.add_heading("第一章 初出茅廬", level=2)
            doc.add_paragraph("少年邁開步伐。")
            doc.add_heading("第二章 險象環生", level=2)
            doc.add_paragraph("四周一片漆黑。")
            doc.save(docx_path)

            options = ImportOptions(mode="markdown")  # mode != novel_regex 時嘗試 heading
            nodes = ImportService.parse_docx(docx_path, "測試Word", options)
            self.assertEqual(len(nodes), 1)
            self.assertEqual(nodes[0].name, "第一卷 風起雲湧")
            self.assertEqual(nodes[0].node_type, "folder")
            self.assertEqual(len(nodes[0].children), 2)
            self.assertEqual(nodes[0].children[0].name, "第一章 初出茅廬")
            self.assertIn("少年邁開步伐", nodes[0].children[0].content)

    def test_directory_parsing(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            # 建立目錄：
            # tmpdir/
            #   ├── 01_第一卷/
            #   │   └── 01_第一章.txt
            #   └── 02_第二卷/
            #       └── 01_第二章.txt
            vol1_dir = os.path.join(tmpdir, "01_第一卷")
            vol2_dir = os.path.join(tmpdir, "02_第二卷")
            os.makedirs(vol1_dir)
            os.makedirs(vol2_dir)

            with open(os.path.join(vol1_dir, "01_第一章.txt"), "w", encoding="utf-8") as f:
                f.write("這是第一章的內文。")
            with open(os.path.join(vol2_dir, "01_第二章.txt"), "w", encoding="utf-8") as f:
                f.write("這是第二章的內文。")

            options = ImportOptions(mode="single_chapter")
            nodes = ImportService.parse_directory(tmpdir, options)
            self.assertEqual(len(nodes), 2)
            self.assertEqual(nodes[0].name, "01_第一卷")
            self.assertEqual(nodes[0].node_type, "folder")
            self.assertEqual(len(nodes[0].children), 1)
            self.assertEqual(nodes[0].children[0].name, "01_第一章")
            self.assertEqual(nodes[0].children[0].content, "這是第一章的內文。")

if __name__ == "__main__":
    unittest.main()
