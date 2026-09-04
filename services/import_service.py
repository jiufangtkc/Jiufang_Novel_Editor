import os
import re
import uuid
from dataclasses import dataclass, field
from typing import List, Optional, Tuple, Literal
from models.models import ChapterNode

# 預設中文小說正則表達式
DEFAULT_VOLUME_REGEX = r"^[ \t]*(?:第[一二三四五六七八九十百千0-9]+[卷部篇]|卷[一二三四五六七八九十0-9]+)[ \t]*(.*)$"
DEFAULT_CHAPTER_REGEX = r"^[ \t]*(?:第[一二三四五六七八九十百千0-9]+[章回節]|Chapter\s*\d+|[0-9]+[\.、\s])[ \t]*(.*)$"
DEFAULT_SCENE_REGEX = r"^[ \t]*(?:\*{3,}|-{3,}|◆{3,}|={3,}|——{2,}|──{2,})[ \t]*$"

@dataclass
class ImportOptions:
    mode: Literal["novel_regex", "markdown", "single_chapter", "custom_regex"] = "novel_regex"
    volume_regex: str = DEFAULT_VOLUME_REGEX
    chapter_regex: str = DEFAULT_CHAPTER_REGEX
    enable_scene_split: bool = False
    scene_regex: str = DEFAULT_SCENE_REGEX
    encoding: str = "auto"  # "auto", "utf-8", "cp950", "gb18030", 等
    prologue_name: str = "序言 / 前言"
    default_chapter_name: str = "未命名章節"


class ImportService:
    """專門負責將外部檔案（.txt, .md, .docx）或資料夾解析並轉換為 ChapterNode 樹狀結構的服務。"""

    @staticmethod
    def detect_encoding(file_path: str) -> str:
        """偵測檔案編碼。優先測試 UTF-8 / UTF-8-SIG，若失敗則對比 CP950 與 GB18030 的中文字可讀性。"""
        with open(file_path, "rb") as f:
            raw = f.read(16384)

        if not raw:
            return "utf-8"

        # 1. 檢查 UTF-8 BOM
        if raw.startswith(b"\xef\xbb\xbf"):
            return "utf-8-sig"

        # 2. 嚴格測試 UTF-8
        try:
            raw.decode("utf-8")
            return "utf-8"
        except UnicodeDecodeError:
            pass

        # 3. 評估 CP950 (Big5) 與 GB18030 (簡體中文)
        def score_decoded(text: str) -> int:
            score = 0
            common_chars = "，。！？「」『』、的一是在不了有和人這中大為上個國我以要他時來用們生到作地於出就分對成會可主發年動同"
            for c in common_chars:
                score += text.count(c)
            # 扣分：私用區或過多難字
            for ch in text:
                code = ord(ch)
                if 0xE000 <= code <= 0xF8FF:  # 私用區 (PUA) 經常是解錯碼的亂碼
                    score -= 5
            return score

        score_cp950 = -99999
        try:
            t_cp950 = raw.decode("cp950")
            score_cp950 = score_decoded(t_cp950)
        except Exception:
            pass

        score_gb = -99999
        try:
            t_gb = raw.decode("gb18030")
            score_gb = score_decoded(t_gb)
        except Exception:
            pass

        if score_cp950 > score_gb:
            return "cp950"
        elif score_gb > score_cp950 and score_gb > 0:
            return "gb18030"

        # 預設回退
        return "utf-8"

    @staticmethod
    def read_file_content(file_path: str, specified_encoding: str = "auto") -> str:
        """讀取純文字或 Markdown 檔案內容，可指定編碼或自動偵測。"""
        if not specified_encoding or specified_encoding.lower() == "auto":
            enc = ImportService.detect_encoding(file_path)
        else:
            enc = specified_encoding
        try:
            with open(file_path, "r", encoding=enc, errors="replace") as f:
                return f.read()
        except Exception:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()

    @classmethod
    def parse_file(cls, file_path: str, options: Optional[ImportOptions] = None) -> List[ChapterNode]:
        """統一解析單一檔案（支援 .txt, .md, .docx）。"""
        if options is None:
            options = ImportOptions()

        ext = os.path.splitext(file_path)[1].lower()
        base_name = os.path.splitext(os.path.basename(file_path))[0]

        if ext == ".docx":
            return cls.parse_docx(file_path, base_name, options)
        elif ext == ".md" and options.mode == "markdown":
            content = cls.read_file_content(file_path, options.encoding)
            return cls.parse_markdown(content, base_name)
        else:
            # .txt 或將 .md 依常規小說 regex 解析
            content = cls.read_file_content(file_path, options.encoding)
            return cls.parse_text(content, base_name, options)

    @classmethod
    def parse_text(cls, text: str, default_title: str, options: ImportOptions) -> List[ChapterNode]:
        """核心純文字解析狀態機：將大篇幅文字切割為 卷(folder) -> 章(file) -> 場景(scene)。"""
        if options.mode == "single_chapter":
            return [
                ChapterNode(
                    name=default_title or options.default_chapter_name,
                    node_type="file",
                    content=text.strip()
                )
            ]

        # 編譯 regex
        vol_pattern = None
        if options.volume_regex:
            try:
                vol_pattern = re.compile(options.volume_regex, re.IGNORECASE)
            except re.error:
                vol_pattern = None

        chap_pattern = None
        if options.chapter_regex:
            try:
                chap_pattern = re.compile(options.chapter_regex, re.IGNORECASE)
            except re.error:
                chap_pattern = None

        scene_pattern = None
        if options.enable_scene_split and options.scene_regex:
            try:
                scene_pattern = re.compile(options.scene_regex)
            except re.error:
                scene_pattern = None

        lines = text.splitlines()
        root_nodes: List[ChapterNode] = []

        current_volume: Optional[ChapterNode] = None
        current_chapter: Optional[ChapterNode] = None
        current_scene: Optional[ChapterNode] = None

        buffer_lines: List[str] = []

        def flush_buffer():
            nonlocal buffer_lines
            content = "\n".join(buffer_lines).strip()
            buffer_lines = []
            if current_scene is not None:
                current_scene.content = content
            elif current_chapter is not None:
                current_chapter.content = content
            elif content:
                # 在遇到第一個章節或分卷前的文字，歸類為前言/序言
                prologue_node = ChapterNode(
                    name=options.prologue_name,
                    node_type="file",
                    content=content
                )
                if current_volume is not None:
                    current_volume.children.append(prologue_node)
                else:
                    root_nodes.append(prologue_node)

        for line in lines:
            stripped = line.strip()
            if not stripped:
                buffer_lines.append(line)
                continue

            # 1. 檢查是否符合分卷標題
            if vol_pattern and vol_pattern.match(stripped):
                flush_buffer()
                current_scene = None
                current_chapter = None
                current_volume = ChapterNode(
                    name=stripped,
                    node_type="folder",
                    children=[]
                )
                root_nodes.append(current_volume)
                continue

            # 2. 檢查是否符合章節標題
            if chap_pattern and chap_pattern.match(stripped):
                flush_buffer()
                current_scene = None
                current_chapter = ChapterNode(
                    name=stripped,
                    node_type="file",
                    content=""
                )
                if current_volume is not None:
                    current_volume.children.append(current_chapter)
                else:
                    root_nodes.append(current_chapter)
                continue

            # 3. 檢查是否符合場景切分線
            if scene_pattern and scene_pattern.match(stripped) and current_chapter is not None:
                flush_buffer()
                scene_count = len(current_chapter.children) + 1
                current_scene = ChapterNode(
                    name=f"場景 {scene_count}",
                    node_type="scene",
                    content=""
                )
                current_chapter.children.append(current_scene)
                continue

            buffer_lines.append(line)

        flush_buffer()

        # 若全文未匹配到任何章節，則整篇做為單一章節傳回
        if not root_nodes:
            root_nodes.append(
                ChapterNode(
                    name=default_title or options.default_chapter_name,
                    node_type="file",
                    content=text.strip()
                )
            )

        return root_nodes

    @classmethod
    def parse_markdown(cls, markdown_text: str, default_title: str) -> List[ChapterNode]:
        """依 Markdown 標題層級 (#, ##, ###) 解析為 卷 -> 章 -> 場景 樹狀結構。"""
        lines = markdown_text.splitlines()
        root_nodes: List[ChapterNode] = []

        current_volume: Optional[ChapterNode] = None
        current_chapter: Optional[ChapterNode] = None
        current_scene: Optional[ChapterNode] = None
        buffer_lines: List[str] = []

        def flush_buffer():
            nonlocal buffer_lines
            content = "\n".join(buffer_lines).strip()
            buffer_lines = []
            if current_scene is not None:
                current_scene.content = content
            elif current_chapter is not None:
                current_chapter.content = content
            elif content:
                prologue_node = ChapterNode(
                    name="序言 / 導言",
                    node_type="file",
                    content=content
                )
                if current_volume is not None:
                    current_volume.children.append(prologue_node)
                else:
                    root_nodes.append(prologue_node)

        heading_regex = re.compile(r"^(#{1,3})\s+(.+)$")

        for line in lines:
            m = heading_regex.match(line.strip())
            if m:
                level = len(m.group(1))
                title = m.group(2).strip()
                flush_buffer()

                if level == 1:
                    # 一級標題 -> 卷 (Folder)
                    current_scene = None
                    current_chapter = None
                    current_volume = ChapterNode(name=title, node_type="folder", children=[])
                    root_nodes.append(current_volume)
                elif level == 2:
                    # 二級標題 -> 章 (File)
                    current_scene = None
                    current_chapter = ChapterNode(name=title, node_type="file", content="")
                    if current_volume is not None:
                        current_volume.children.append(current_chapter)
                    else:
                        root_nodes.append(current_chapter)
                elif level == 3:
                    # 三級標題 -> 場景 (Scene)
                    current_scene = ChapterNode(name=title, node_type="scene", content="")
                    if current_chapter is not None:
                        current_chapter.children.append(current_scene)
                    elif current_volume is not None:
                        # 若無當前章，自動建一個父章節
                        current_chapter = ChapterNode(name="未命名章節", node_type="file", content="")
                        current_volume.children.append(current_chapter)
                        current_chapter.children.append(current_scene)
                    else:
                        current_chapter = ChapterNode(name="未命名章節", node_type="file", content="")
                        root_nodes.append(current_chapter)
                        current_chapter.children.append(current_scene)
                continue

            buffer_lines.append(line)

        flush_buffer()

        if not root_nodes:
            root_nodes.append(
                ChapterNode(
                    name=default_title or "未命名 Markdown",
                    node_type="file",
                    content=markdown_text.strip()
                )
            )

        return root_nodes

    @classmethod
    def parse_docx(cls, docx_path: str, default_title: str, options: ImportOptions) -> List[ChapterNode]:
        """讀取 Word .docx 文件，支援 Heading 樣式或中文小說標題正則表達式。"""
        try:
            from docx import Document
        except ImportError:
            raise RuntimeError("系統環境未安裝 python-docx，無法解析 Word 文件。")

        doc = Document(docx_path)
        # 檢查段落中是否有 Heading 1 / Heading 2
        has_headings = False
        for p in doc.paragraphs:
            if p.style and p.style.name in ("Heading 1", "Heading 2", "Heading 3", "標題 1", "標題 2", "標題 3"):
                has_headings = True
                break

        if has_headings and options.mode != "novel_regex":
            # 使用大綱樣式解析
            root_nodes: List[ChapterNode] = []
            current_volume: Optional[ChapterNode] = None
            current_chapter: Optional[ChapterNode] = None
            current_scene: Optional[ChapterNode] = None
            buffer_lines: List[str] = []

            def flush_buffer():
                nonlocal buffer_lines
                content = "\n".join(buffer_lines).strip()
                buffer_lines = []
                if current_scene is not None:
                    current_scene.content = content
                elif current_chapter is not None:
                    current_chapter.content = content
                elif content:
                    node = ChapterNode(name=options.prologue_name, node_type="file", content=content)
                    if current_volume:
                        current_volume.children.append(node)
                    else:
                        root_nodes.append(node)

            for p in doc.paragraphs:
                text = p.text.strip()
                style_name = p.style.name if p.style else ""

                if style_name in ("Heading 1", "標題 1"):
                    flush_buffer()
                    current_scene = None
                    current_chapter = None
                    current_volume = ChapterNode(name=text or "未命名分卷", node_type="folder", children=[])
                    root_nodes.append(current_volume)
                elif style_name in ("Heading 2", "標題 2"):
                    flush_buffer()
                    current_scene = None
                    current_chapter = ChapterNode(name=text or "未命名章節", node_type="file", content="")
                    if current_volume:
                        current_volume.children.append(current_chapter)
                    else:
                        root_nodes.append(current_chapter)
                elif style_name in ("Heading 3", "標題 3"):
                    flush_buffer()
                    current_scene = ChapterNode(name=text or "場景", node_type="scene", content="")
                    if current_chapter:
                        current_chapter.children.append(current_scene)
                    else:
                        current_chapter = ChapterNode(name="未命名章節", node_type="file", content="")
                        root_nodes.append(current_chapter)
                        current_chapter.children.append(current_scene)
                else:
                    if text:
                        buffer_lines.append(p.text)
                    else:
                        buffer_lines.append("")

            flush_buffer()
            if not root_nodes:
                full_text = "\n".join(p.text for p in doc.paragraphs)
                root_nodes.append(ChapterNode(name=default_title, node_type="file", content=full_text.strip()))
            return root_nodes
        else:
            # 將全部段落組合成純文字，使用文字正規表示式解析
            full_text = "\n".join(p.text for p in doc.paragraphs)
            return cls.parse_text(full_text, default_title, options)

    @classmethod
    def parse_directory(cls, dir_path: str, options: Optional[ImportOptions] = None) -> List[ChapterNode]:
        """批次掃描資料夾：子資料夾映射為 Folder，.txt / .md / .docx 映射為 File。"""
        if options is None:
            options = ImportOptions()

        try:
            import natsort
            sort_fn = natsort.natsorted
        except ImportError:
            sort_fn = sorted

        def scan_dir(path: str) -> List[ChapterNode]:
            nodes = []
            try:
                entries = sort_fn(os.listdir(path))
            except Exception:
                return []

            for entry in entries:
                full_path = os.path.join(path, entry)
                if os.path.isdir(full_path):
                    sub_children = scan_dir(full_path)
                    folder_node = ChapterNode(
                        name=entry,
                        node_type="folder",
                        children=sub_children
                    )
                    nodes.append(folder_node)
                elif os.path.isfile(full_path):
                    ext = os.path.splitext(entry)[1].lower()
                    if ext in (".txt", ".md", ".docx"):
                        base_name = os.path.splitext(entry)[0]
                        file_nodes = cls.parse_file(full_path, options)
                        if len(file_nodes) == 1:
                            file_nodes[0].name = base_name
                            nodes.append(file_nodes[0])
                        else:
                            file_folder = ChapterNode(
                                name=base_name,
                                node_type="folder",
                                children=file_nodes
                            )
                            nodes.append(file_folder)
            return nodes

        return scan_dir(dir_path)
