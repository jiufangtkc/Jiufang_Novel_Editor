import re
import html
from typing import List, Tuple, Dict, Any


class MarkdownConverter:
    """
    小說 Markdown 結構化轉換中介器。
    負責將純 Markdown 文本解析並轉換為：
    1. 乾淨小說純文字 (去除語法符號，保留或套用全形縮排)
    2. HTML / XHTML 標籤 (供 ePub 電子書與 Rich Text 使用)
    3. Docx Paragraphs & Runs 結構 (供 Word 匯出使用)
    """

    # 行內格式正則表示式 (粗斜體、粗體、斜體、刪除線、行內代碼)
    # 優先匹配最長語法
    INLINE_PATTERN = re.compile(
        r'(?P<bold_italic>\*\*\*(?P<bi_text>[^\*\n]+?)\*\*\*)|'
        r'(?P<bold>\*\*(?P<b_text>[^\*\n]+?)\*\*|__(?P<b2_text>[^_\n]+?)__)|'
        r'(?P<italic>\*(?P<i_text>[^\*\n]+?)\*|_(?P<i2_text>[^_\n]+?)_)|'
        r'(?P<strike>~~(?P<s_text>[^~\n]+?)~~)|'
        r'(?P<code>`(?P<c_text>[^`\n]+?)`)'
    )

    @classmethod
    def parse_inline_tokens(cls, text: str) -> List[Dict[str, Any]]:
        """
        將單行文本解析為 Token 列表。
        每個 Token 包含:
        - text: 實際文字內容
        - bold: bool
        - italic: bool
        - strike: bool
        - code: bool
        """
        tokens = []
        last_idx = 0

        for match in cls.INLINE_PATTERN.finditer(text):
            start, end = match.span()
            if start > last_idx:
                tokens.append({
                    "text": text[last_idx:start],
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "code": False
                })

            groupdict = match.groupdict()
            if groupdict.get("bold_italic"):
                tokens.append({
                    "text": groupdict["bi_text"],
                    "bold": True,
                    "italic": True,
                    "strike": False,
                    "code": False
                })
            elif groupdict.get("bold"):
                tokens.append({
                    "text": groupdict.get("b_text") or groupdict.get("b2_text"),
                    "bold": True,
                    "italic": False,
                    "strike": False,
                    "code": False
                })
            elif groupdict.get("italic"):
                tokens.append({
                    "text": groupdict.get("i_text") or groupdict.get("i2_text"),
                    "bold": False,
                    "italic": True,
                    "strike": False,
                    "code": False
                })
            elif groupdict.get("strike"):
                tokens.append({
                    "text": groupdict["s_text"],
                    "bold": False,
                    "italic": False,
                    "strike": True,
                    "code": False
                })
            elif groupdict.get("code"):
                tokens.append({
                    "text": groupdict["c_text"],
                    "bold": False,
                    "italic": False,
                    "strike": False,
                    "code": True
                })

            last_idx = end

        if last_idx < len(text):
            tokens.append({
                "text": text[last_idx:],
                "bold": False,
                "italic": False,
                "strike": False,
                "code": False
            })

        return tokens

    @classmethod
    def to_plain_text(cls, md_text: str, auto_indent: bool = True) -> str:
        """
        將 Markdown 文字轉換為乾淨純文字，去除 Markdown 標記語法。
        如果 auto_indent 為 True，則為有內容的段落補上全形縮排兩格 (　　)。
        """
        if not md_text:
            return ""

        lines = md_text.split('\n')
        out_lines = []

        for line in lines:
            stripped = line.strip()

            # 處理分隔線
            if re.match(r'^(?:---|\*\*\*|___)$', stripped):
                out_lines.append("――――――――――――――――――――")
                continue

            # 處理標題 (去除 # 開頭)
            header_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if header_match:
                header_text = header_match.group(2).strip()
                # 去除行內標記
                tokens = cls.parse_inline_tokens(header_text)
                clean_header = "".join(t["text"] for t in tokens)
                out_lines.append(clean_header)
                continue

            # 處理引言 (去除 > 開頭)
            quote_match = re.match(r'^>\s*(.*)$', stripped)
            if quote_match:
                q_text = quote_match.group(1).strip()
                tokens = cls.parse_inline_tokens(q_text)
                clean_q = "".join(t["text"] for t in tokens)
                if auto_indent and clean_q:
                    out_lines.append(f"　　「{clean_q}」" if not clean_q.startswith("「") else f"　　{clean_q}")
                else:
                    out_lines.append(clean_q)
                continue

            # 一般段落
            if stripped:
                tokens = cls.parse_inline_tokens(line)
                clean_text = "".join(t["text"] for t in tokens).strip()
                if auto_indent:
                    if not clean_text.startswith("　　") and not clean_text.startswith("    "):
                        out_lines.append(f"　　{clean_text}")
                    else:
                        out_lines.append(clean_text)
                else:
                    out_lines.append(clean_text)
            else:
                out_lines.append("")

        return "\n".join(out_lines)

    @classmethod
    def to_html_paragraphs(cls, md_text: str) -> List[str]:
        """
        將 Markdown 文本轉換為 HTML 標籤段落列表。
        支援標題、分隔線、引言、行內樣式 (<strong>, <em>, <del>, <code>)。
        """
        if not md_text:
            return []

        lines = md_text.split('\n')
        html_paragraphs = []

        for line in lines:
            stripped = line.strip()

            # 空行
            if not stripped:
                html_paragraphs.append("<p><br/></p>")
                continue

            # 分隔線
            if re.match(r'^(?:---|\*\*\*|___)$', stripped):
                html_paragraphs.append("<hr/>")
                continue

            # 標題
            header_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if header_match:
                level = len(header_match.group(1))
                h_text = header_match.group(2).strip()
                inner_html = cls._tokens_to_html(cls.parse_inline_tokens(h_text))
                html_paragraphs.append(f"<h{level}>{inner_html}</h{level}>")
                continue

            # 引言
            quote_match = re.match(r'^>\s*(.*)$', stripped)
            if quote_match:
                q_text = quote_match.group(1).strip()
                inner_html = cls._tokens_to_html(cls.parse_inline_tokens(q_text))
                html_paragraphs.append(f"<blockquote><p>{inner_html}</p></blockquote>")
                continue

            # 一般小說段落
            inner_html = cls._tokens_to_html(cls.parse_inline_tokens(line.rstrip()))
            html_paragraphs.append(f"<p>{inner_html}</p>")

        return html_paragraphs

    @classmethod
    def _tokens_to_html(cls, tokens: List[Dict[str, Any]]) -> str:
        res = []
        for t in tokens:
            piece = html.escape(t["text"])
            if t["code"]:
                piece = f"<code>{piece}</code>"
            if t["strike"]:
                piece = f"<del>{piece}</del>"
            if t["italic"]:
                piece = f"<em>{piece}</em>"
            if t["bold"]:
                piece = f"<strong>{piece}</strong>"
            res.append(piece)
        return "".join(res)

    @classmethod
    def render_to_docx(cls, md_text: str, doc, font_family: str = "Iansui"):
        """
        將 Markdown 文本渲染至 python-docx Document。
        處理粗體、斜體、刪除線、標題、分隔線、引言及首行縮排。
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not md_text:
            return

        lines = md_text.split('\n')
        for line in lines:
            stripped = line.strip()

            # 1. 空行
            if not stripped:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                continue

            # 2. 分隔線 (--- 或 ***)
            if re.match(r'^(?:---|\*\*\*|___)$', stripped):
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                run = p.add_run("――――――――――――――――――――")
                run.font.name = font_family
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
                continue

            # 3. 標題
            header_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
            if header_match:
                level = len(header_match.group(1))
                h_text = header_match.group(2).strip()
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True
                
                # 標題字級
                size_map = {1: 16, 2: 14, 3: 13, 4: 12, 5: 12, 6: 12}
                tokens = cls.parse_inline_tokens(h_text)
                for t in tokens:
                    run = p.add_run(t["text"])
                    run.bold = True
                    run.italic = t["italic"]
                    run.font.name = font_family
                    run.font.size = Pt(size_map.get(level, 12))
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
                continue

            # 4. 引言
            quote_match = re.match(r'^>\s*(.*)$', stripped)
            if quote_match:
                q_text = quote_match.group(1).strip()
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(24)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.35
                tokens = cls.parse_inline_tokens(q_text)
                for t in tokens:
                    run = p.add_run(t["text"])
                    run.bold = t["bold"]
                    run.italic = True
                    run.font.strike = t["strike"]
                    run.font.name = font_family
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)
                continue

            # 5. 一般小說段落
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.35
            p.paragraph_format.first_line_indent = Pt(24)

            tokens = cls.parse_inline_tokens(line.rstrip())
            for t in tokens:
                run = p.add_run(t["text"])
                run.bold = t["bold"]
                run.italic = t["italic"]
                run.font.strike = t["strike"]
                run.font.name = font_family
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)

