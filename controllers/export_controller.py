import os
import re
import datetime
import html
import zipfile
import uuid
from PyQt6.QtWidgets import QDialog, QMessageBox, QFileDialog
from PyQt6.QtGui import QTextDocument
from PyQt6.QtCore import Qt
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn

from views.dialogs.export_scope_dialog import ExportScopeDialog
from utils.markdown_converter import MarkdownConverter


class ExportController:
    """負責專案與文件匯出之控制器，支援 docx, txt, md, epub 多種格式。"""

    def __init__(self, main_controller):
        self.mc = main_controller
        self.view = main_controller.view

    def export_documents(self, item=None):
        """開啟匯出對話框並執行多格式匯出。"""
        self.mc.save_current_editor_content()
        checked_item = item
        if checked_item is None:
            checked_item = self.mc.current_file_item
        if not self.mc.tree.is_item_valid(checked_item):
            checked_item = self.view.tree_widget.currentItem()

        book_title = getattr(self.mc.project_info, "title", "未命名作品") or "未命名作品"
        dialog = ExportScopeDialog(self.view, checked_item=checked_item, default_title=book_title)
        if dialog.exec() != QDialog.DialogCode.Accepted:
            return

        files_list = dialog.get_checked_files()
        if not files_list:
            QMessageBox.warning(self.view, "提示", "請先在匯出視窗中勾選要匯出的章節。")
            return

        fmt = dialog.get_export_format()  # docx, txt, md, epub
        include_title = dialog.is_include_title()
        merge_mode = dialog.is_merge_mode()

        # 準備預設目錄
        default_dir = self.mc.get_export_dir()
        os.makedirs(default_dir, exist_ok=True)
        now_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        # 決定預設檔名
        if len(files_list) == 1:
            raw_name = files_list[0].text(0).strip()
        else:
            raw_name = book_title
        clean_name = re.sub(r'[\/\\\:\*\?\"\<\>\|]', '_', raw_name)

        filter_map = {
            "docx": "Word 文件 (*.docx)",
            "txt": "純文字檔案 (*.txt)",
            "md": "Markdown 檔案 (*.md)",
            "epub": "ePub 電子書 (*.epub)"
        }

        if merge_mode:
            # 合併匯出為單一檔案：讓使用者選擇具體檔案儲存路徑
            default_path = os.path.join(default_dir, f"{clean_name}_{now_str}.{fmt}")
            save_path, _ = QFileDialog.getSaveFileName(
                self.view,
                "選擇匯出儲存位置",
                default_path,
                filter_map.get(fmt, "所有檔案 (*.*)")
            )
            if not save_path:
                return

            self._export_merged(files_list, save_path, fmt, include_title, book_title)
        else:
            # 分割匯出為多個檔案：讓使用者選擇目標資料夾
            target_dir = QFileDialog.getExistingDirectory(
                self.view,
                "選擇匯出目標資料夾",
                default_dir
            )
            if not target_dir:
                return

            self._export_separated(files_list, target_dir, fmt, include_title, book_title)

    def _export_merged(self, files_list, save_path: str, fmt: str, include_title: bool, book_title: str):
        """合併所有選取章節並匯出為單一檔案。"""
        try:
            if fmt == "docx":
                self._save_as_docx(files_list, save_path, include_title)
            elif fmt == "txt":
                self._save_as_txt(files_list, save_path, include_title)
            elif fmt == "md":
                self._save_as_md(files_list, save_path, include_title)
            elif fmt == "epub":
                self._save_as_epub(files_list, save_path, include_title, book_title)
            else:
                self._save_as_txt(files_list, save_path, include_title)

            QMessageBox.information(
                self.view, "匯出成功",
                f"檔案已成功匯出至：\n{save_path}"
            )
        except Exception as e:
            QMessageBox.critical(self.view, "匯出失敗", f"儲存檔案時發生錯誤：\n{e}")

    def _export_separated(self, files_list, target_dir: str, fmt: str, include_title: bool, book_title: str):
        """將章節分別獨立匯出至目標資料夾。"""
        try:
            count = 0
            for idx, file_item in enumerate(files_list, 1):
                raw_title = file_item.text(0).strip() or f"章節_{idx}"
                clean_title = re.sub(r'[\/\\\:\*\?\"\<\>\|]', '_', raw_title)
                filename = f"{idx:02d}_{clean_title}.{fmt}"
                save_path = os.path.join(target_dir, filename)

                if fmt == "docx":
                    self._save_as_docx([file_item], save_path, include_title)
                elif fmt == "txt":
                    self._save_as_txt([file_item], save_path, include_title)
                elif fmt == "md":
                    self._save_as_md([file_item], save_path, include_title)
                elif fmt == "epub":
                    self._save_as_epub([file_item], save_path, include_title, raw_title)
                count += 1

            QMessageBox.information(
                self.view, "匯出成功",
                f"已成功將 {count} 個章節分別匯出至資料夾：\n{target_dir}"
            )
        except Exception as e:
            QMessageBox.critical(self.view, "匯出失敗", f"批次匯出檔案時發生錯誤：\n{e}")

    # =========================================================================
    # 各格式具體生成實作
    # =========================================================================

    def _get_item_plain_text(self, file_item) -> str:
        node_data = file_item.data(0, Qt.ItemDataRole.UserRole)
        content = node_data.get("content", "") if node_data else ""
        return content

    def _save_as_docx(self, files_list, save_path: str, include_title: bool):
        doc = Document()
        font_family = self.mc.global_font_family or "Iansui"

        normal_style = doc.styles['Normal']
        normal_style.font.name = font_family
        normal_style.font.size = Pt(12)
        normal_style.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)

        for file_item in files_list:
            title = file_item.text(0).strip()
            raw_content = self._get_item_plain_text(file_item)

            if include_title and title:
                heading = doc.add_paragraph()
                heading.paragraph_format.space_before = Pt(14)
                heading.paragraph_format.space_after = Pt(8)
                heading.paragraph_format.keep_with_next = True
                
                run = heading.add_run(title)
                run.bold = True
                run.font.size = Pt(16)
                run.font.name = font_family
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_family)

            if raw_content:
                MarkdownConverter.render_to_docx(raw_content, doc, font_family)

        doc.save(save_path)

    def _save_as_txt(self, files_list, save_path: str, include_title: bool):
        lines = []
        for file_item in files_list:
            title = file_item.text(0).strip()
            raw_content = self._get_item_plain_text(file_item)

            if include_title and title:
                lines.append(f"【{title}】\n")

            if raw_content:
                clean_text = MarkdownConverter.to_plain_text(raw_content, auto_indent=True)
                lines.append(clean_text)
            lines.append("\n")

        with open(save_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(lines))

    def _save_as_md(self, files_list, save_path: str, include_title: bool):
        lines = []
        for file_item in files_list:
            title = file_item.text(0).strip()
            raw_content = self._get_item_plain_text(file_item)

            if include_title and title:
                lines.append(f"# {title}\n")

            if raw_content:
                paragraphs = raw_content.split('\n')
                for p in paragraphs:
                    lines.append(p.rstrip())
            lines.append("\n---\n")

        with open(save_path, "w", encoding="utf-8") as f:
            f.write('\n'.join(lines))

    def _save_as_epub(self, files_list, save_path: str, include_title: bool, book_title: str):
        book_id = str(uuid.uuid4())
        author = getattr(self.mc.project_info, "author", "作者") if hasattr(self.mc, 'project_info') else "作者"

        with zipfile.ZipFile(save_path, 'w') as zf:
            # 1. mimetype (未壓縮)
            zf.writestr("mimetype", "application/epub+zip", compress_type=zipfile.ZIP_STORED)

            # 2. META-INF/container.xml
            container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
  <rootfiles>
    <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
  </rootfiles>
</container>"""
            zf.writestr("META-INF/container.xml", container_xml, compress_type=zipfile.ZIP_DEFLATED)

            # 3. CSS
            css_content = """
body { font-family: "PingFang TC", "Heiti TC", "Microsoft JhengHei", "Iansui", serif; line-height: 1.7; margin: 4%; }
h1 { text-align: center; margin-top: 1.5em; margin-bottom: 1.2em; font-size: 1.5em; color: #222; }
h2 { margin-top: 1.2em; margin-bottom: 0.8em; font-size: 1.3em; color: #333; }
h3 { margin-top: 1.0em; margin-bottom: 0.6em; font-size: 1.1em; color: #444; }
p { text-indent: 2em; margin-top: 0.4em; margin-bottom: 0.4em; text-align: justify; }
blockquote { margin: 1em 2em; color: #555; border-left: 3px solid #ccc; padding-left: 1em; }
hr { border: none; border-top: 1px dashed #aaa; margin: 2em auto; width: 60%; }
"""
            zf.writestr("OEBPS/style.css", css_content, compress_type=zipfile.ZIP_DEFLATED)

            # 4. 章節 XHTML
            manifest_items = ['<item id="style" href="style.css" media-type="text/css"/>']
            spine_items = []
            toc_navpoints = []

            for idx, file_item in enumerate(files_list, 1):
                ch_title = file_item.text(0).strip() or f"第 {idx} 章"
                ch_content = self._get_item_plain_text(file_item)

                ch_filename = f"chapter_{idx}.xhtml"
                ch_id = f"chapter_{idx}"

                manifest_items.append(f'<item id="{ch_id}" href="{ch_filename}" media-type="application/xhtml+xml"/>')
                spine_items.append(f'<itemref idref="{ch_id}"/>')

                toc_navpoints.append(f"""  <navPoint id="navPoint-{idx}" playOrder="{idx}">
    <navLabel><text>{html.escape(ch_title)}</text></navLabel>
    <content src="{ch_filename}"/>
  </navPoint>""")

                p_tags = MarkdownConverter.to_html_paragraphs(ch_content)
                title_html = f"<h1>{html.escape(ch_title)}</h1>" if include_title else ""
                ch_html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
  <title>{html.escape(ch_title)}</title>
  <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
  {title_html}
  {''.join(p_tags)}
</body>
</html>"""
                zf.writestr(f"OEBPS/{ch_filename}", ch_html, compress_type=zipfile.ZIP_DEFLATED)

            # 5. toc.ncx
            ncx_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<ncx xmlns="http://www.daisy.org/z3986/2005/ncx/" version="2005-1">
  <head>
    <meta name="dtb:uid" content="{book_id}"/>
    <meta name="dtb:depth" content="1"/>
    <meta name="dtb:totalPageCount" content="0"/>
    <meta name="dtb:maxPageNumber" content="0"/>
  </head>
  <docTitle><text>{html.escape(book_title)}</text></docTitle>
  <navMap>
{''.join(toc_navpoints)}
  </navMap>
</ncx>"""
            zf.writestr("OEBPS/toc.ncx", ncx_content, compress_type=zipfile.ZIP_DEFLATED)
            manifest_items.append('<item id="ncx" href="toc.ncx" media-type="application/x-dtbncx+xml"/>')

            # 6. content.opf
            manifest_str = '\n    '.join(manifest_items)
            spine_str = '\n    '.join(spine_items)
            opf_content = f"""<?xml version="1.0" encoding="utf-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="BookId" version="2.0">
  <metadata xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:opf="http://www.idpf.org/2007/opf">
    <dc:title>{html.escape(book_title)}</dc:title>
    <dc:language>zh-TW</dc:language>
    <dc:identifier id="BookId" opf:scheme="UUID">{book_id}</dc:identifier>
    <dc:creator>{html.escape(author)}</dc:creator>
  </metadata>
  <manifest>
    {manifest_str}
  </manifest>
  <spine toc="ncx">
    {spine_str}
  </spine>
</package>"""
            zf.writestr("OEBPS/content.opf", opf_content, compress_type=zipfile.ZIP_DEFLATED)
